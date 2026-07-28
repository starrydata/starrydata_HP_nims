#!/usr/bin/env python3
"""
Starrydata2 マニュアル Word (.docx) ビルダー
src/manual/*.njk を読み、編集可能な Word 文書を生成する。
出力: docs/Starrydata2_Manual_ja.docx
"""

import re
from pathlib import Path
from bs4 import BeautifulSoup, NavigableString
from docx import Document
from docx.shared import Pt, Mm, Cm, RGBColor, Emu, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image as PILImage


REPO = Path("/Users/atsumitanaka/Documents/starrydata_HP")
SRC_DIR = REPO / "src" / "manual"
OUT = REPO / "docs" / "Starrydata2_Manual_ja.docx"

INK = RGBColor(0x0F, 0x17, 0x2A)
MUTED = RGBColor(0x64, 0x74, 0x8B)
ACCENT = RGBColor(0x4F, 0x46, 0xE5)
WARN_FILL = "FFF7E6"
INFO_FILL = "E8F5E9"
CODE_FILL = "0D1B3E"
CODE_INK = RGBColor(0xE2, 0xE8, 0xF0)
HEAD_FILL = "EEF2FF"


# ---- njk からプレーンな HTML を抜き出す ----
def clean_njk(text: str) -> str:
    text = re.sub(r"^---\n.*?\n---\n", "", text, count=1, flags=re.DOTALL)
    text = re.sub(r"\{%-?\s*set\s+[^%]+%\}", "", text)
    text = re.sub(r"\{\{[^}]+\}\}", "", text)
    return text


# ---- セル背景色 ----
def set_cell_shading(cell, fill_hex: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def set_cell_left_border(cell, color_hex: str, width=24):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(width))
    left.set(qn("w:color"), color_hex)
    borders.append(left)
    tc_pr.append(borders)


# ---- インラインを Run に展開 ----
def add_runs(paragraph, node, *, bold=False, italic=False, mono=False, color=None, underline=False):
    if isinstance(node, NavigableString):
        text = str(node).replace(" ", " ")
        if not text:
            return
        run = paragraph.add_run(text)
        run.bold = bold
        run.italic = italic
        if mono:
            run.font.name = "Menlo"
        if color is not None:
            run.font.color.rgb = color
        if underline:
            run.underline = True
        return

    name = node.name
    if name in ("strong", "b"):
        for c in node.children:
            add_runs(paragraph, c, bold=True, italic=italic, mono=mono, color=color, underline=underline)
        return
    if name in ("em", "i"):
        for c in node.children:
            add_runs(paragraph, c, bold=bold, italic=True, mono=mono, color=color, underline=underline)
        return
    if name == "code":
        for c in node.children:
            add_runs(paragraph, c, bold=bold, italic=italic, mono=True, color=color, underline=underline)
        return
    if name == "a":
        for c in node.children:
            add_runs(paragraph, c, bold=bold, italic=italic, mono=mono,
                     color=ACCENT, underline=True)
        href = node.get("href", "")
        if href:
            r = paragraph.add_run(f" ({href})")
            r.font.size = Pt(8)
            r.font.color.rgb = MUTED
        return
    if name == "br":
        paragraph.add_run().add_break()
        return

    for c in node.children:
        add_runs(paragraph, c, bold=bold, italic=italic, mono=mono, color=color, underline=underline)


# ---- ブロックを Document に書き出す ----
def render(node, doc, content_width_emu):
    name = getattr(node, "name", None)
    if name is None:
        return
    if name in ("script", "style"):
        return

    classes = node.get("class") or []
    style_attr = (node.get("style") or "").replace(" ", "")

    # eyebrow
    if "eyebrow" in classes:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(node.get_text(strip=True))
        r.font.size = Pt(9)
        r.font.color.rgb = MUTED
        r.font.name = "Helvetica"
        return

    if name == "h1":
        p = doc.add_heading(level=0)
        add_runs(p, node)
        return

    if name == "h2":
        p = doc.add_heading(level=1)
        add_runs(p, node)
        return

    if name == "h3":
        p = doc.add_heading(level=2)
        add_runs(p, node)
        return

    if name == "p":
        p = doc.add_paragraph()
        if "lead" in classes:
            p.paragraph_format.space_after = Pt(10)
            for c in node.children:
                add_runs(p, c)
            for r in p.runs:
                r.font.size = Pt(11)
        else:
            for c in node.children:
                add_runs(p, c)
        return

    if name == "ul":
        for li in node.find_all("li", recursive=False):
            p = doc.add_paragraph(style="List Bullet")
            for c in li.children:
                add_runs(p, c)
        return

    if name == "ol":
        for li in node.find_all("li", recursive=False):
            p = doc.add_paragraph(style="List Number")
            for c in li.children:
                add_runs(p, c)
        return

    if name == "pre":
        code_text = node.get_text("\n")
        t = doc.add_table(rows=1, cols=1)
        t.autofit = False
        cell = t.cell(0, 0)
        cell.width = Emu(content_width_emu)
        set_cell_shading(cell, CODE_FILL)
        cell.text = ""
        for line in code_text.split("\n"):
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(line if line else " ")
            r.font.name = "Menlo"
            r.font.size = Pt(9)
            r.font.color.rgb = CODE_INK
        # remove the initial empty paragraph
        first = cell.paragraphs[0]
        if not first.text:
            first._element.getparent().remove(first._element)
        doc.add_paragraph()
        return

    if name == "table":
        rows = []
        for tr in node.find_all("tr"):
            cells = tr.find_all(["th", "td"])
            rows.append([(c.name, c) for c in cells])
        if not rows:
            return
        ncols = max(len(r) for r in rows)
        t = doc.add_table(rows=len(rows), cols=ncols)
        t.style = "Table Grid"
        col_w = Emu(content_width_emu // ncols)
        for i, row in enumerate(rows):
            for j, (cname, cnode) in enumerate(row):
                cell = t.rows[i].cells[j]
                cell.width = col_w
                cell.text = ""
                if cname == "th":
                    set_cell_shading(cell, HEAD_FILL)
                p = cell.paragraphs[0]
                for c in cnode.children:
                    add_runs(p, c, bold=(cname == "th"))
                for r in p.runs:
                    r.font.size = Pt(9)
        doc.add_paragraph()
        return

    # 警告 / 情報ボックス
    if name == "div" and "border-left:4pxsolid" in style_attr:
        bg = WARN_FILL if "f0b429" in style_attr else INFO_FILL
        bar_color = "F0B429" if "f0b429" in style_attr else "43A047"
        t = doc.add_table(rows=1, cols=1)
        t.autofit = False
        cell = t.cell(0, 0)
        cell.width = Emu(content_width_emu)
        set_cell_shading(cell, bg)
        set_cell_left_border(cell, bar_color, width=24)
        cell.text = ""
        first = True
        for child in node.children:
            if isinstance(child, NavigableString):
                txt = str(child).strip()
                if not txt:
                    continue
                p = cell.add_paragraph()
                p.add_run(txt)
                first = False
                continue
            if getattr(child, "name", None) == "br":
                continue
            p = cell.add_paragraph()
            for c in child.children:
                add_runs(p, c)
            first = False
        empty = cell.paragraphs[0]
        if not empty.text:
            empty._element.getparent().remove(empty._element)
        doc.add_paragraph()
        return

    # figure / img
    if name == "figure":
        img = node.find("img")
        cap = node.find("figcaption")
        if img is not None:
            add_image(doc, img.get("src", ""), content_width_emu)
        if cap is not None:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(12)
            for c in cap.children:
                add_runs(p, c, italic=True, color=MUTED)
            for r in p.runs:
                r.font.size = Pt(9)
        return

    if name == "img":
        add_image(doc, node.get("src", ""), content_width_emu)
        return

    # コンテナはネスト走査
    for child in node.children:
        if isinstance(child, NavigableString):
            continue
        render(child, doc, content_width_emu)


def add_image(doc, src: str, content_width_emu: int):
    if not src:
        return
    p_local = REPO / "src" / src.lstrip("/")
    if not p_local.exists():
        p = doc.add_paragraph(f"[画像なし: {src}]")
        for r in p.runs:
            r.font.color.rgb = MUTED
        return
    # アスペクト維持で content_width の 80% に
    with PILImage.open(p_local) as im:
        iw, ih = im.size
    target_w = int(content_width_emu * 0.82)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(p_local), width=Emu(target_w))


# ---- 1 ファイル分を書き出す ----
def render_file(path: Path, doc, content_width_emu, first=False):
    if not path.exists():
        print(f"WARN missing {path}")
        return
    raw = path.read_text(encoding="utf-8")
    soup = BeautifulSoup(clean_njk(raw), "html.parser")
    if not first:
        doc.add_page_break()
    for sec in soup.find_all("section", recursive=False):
        for child in sec.children:
            if isinstance(child, NavigableString):
                continue
            if child.name == "div" and any(c.startswith("container") for c in (child.get("class") or [])):
                for g in child.children:
                    if isinstance(g, NavigableString):
                        continue
                    render(g, doc, content_width_emu)
            else:
                render(child, doc, content_width_emu)


# ---- メイン ----
def main():
    doc = Document()

    # A4, 余白 20mm
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)

    content_width_emu = section.page_width - section.left_margin - section.right_margin

    # ベースフォント設定（日本語対応）
    style = doc.styles["Normal"]
    style.font.name = "Hiragino Kaku Gothic ProN"
    style.font.size = Pt(10.5)
    style.font.color.rgb = INK
    rpr = style.element.rPr
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), "Hiragino Kaku Gothic ProN")
    rfonts.set(qn("w:ascii"), "Helvetica")
    rfonts.set(qn("w:hAnsi"), "Helvetica")

    files = [
        SRC_DIR / "index.njk",
        SRC_DIR / "account.njk",
        SRC_DIR / "search.njk",
        SRC_DIR / "register.njk",
        SRC_DIR / "download.njk",
    ]
    for i, f in enumerate(files):
        render_file(f, doc, content_width_emu, first=(i == 0))

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"OK -> {OUT}")


if __name__ == "__main__":
    main()
